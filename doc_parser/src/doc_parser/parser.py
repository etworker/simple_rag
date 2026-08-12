"""
文档解析实现

支持: PDF (.pdf), Word (.docx)

核心解决的问题：
1. 页眉/页脚/水印噪声 → 用坐标位置过滤
2. 表格文字混入正文 → 提取正文时排除表格区域
3. 跨页段落截断 → 流式拼接全文后按语义分段
4. 重复内容（页眉页脚） → 统计高频重复行自动识别并过滤
"""

import copy
import os
import re
from collections import Counter
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from loguru import logger

from doc_parser.models import Document, Paragraph, Table

# 默认配置
DEFAULT_CONFIG = {
    "header_margin_pct": 8,
    "footer_margin_pct": 8,
    "repeat_line_threshold_pct": 30,
    "min_paragraph_length": 10,
    "max_paragraph_length": 600,
    "chapter_patterns": [
        r"^(\d+\.\d+\.\d+)\s+(.+)",
        r"^(\d+\.\d+)\s+(.+)",
        r"^(\d+)\s+(.+)",
        r"^第\s*(\d+)\s*[章节]\s*(.+)",
        # 中文数字编号（公文常见格式）
        r"^([一二三四五六七八九十]+)、\s*(.+)",
        r"^（([一二三四五六七八九十]+)）\s*(.+)",
    ],
    # 单纯数字编号 ^\d+ 的标题最大长度（超过则视为正文而非章节标题）
    # 真实章节标题如 "总则"/"日常管理" 通常 ≤20 字，
    # 而 "6 月份，管理局通过..." 等正文行会误匹配
    "single_number_title_max_length": 15,
    # 单纯数字编号的最大值（超过则视为年份等而非章节号）
    "single_number_max_value": 999,
    "noise_patterns": [
        r"^\s*$",
    ],
    # 元数据行剥离：独立成行的版本管理元数据（修订日期/发布日期/独立日期/版次等）
    # 整行匹配时从段落流中剥离（不进正文段落），避免被并入相邻正文造成句子拼接。
    # 通用文档元数据模式，可被用户 config['extract']['noise_line_patterns'] 覆盖。
    "noise_line_patterns": [
        r"^修订日期\s*[：:]\s*\S+$",
        r"^发布日期\s*[：:]\s*\S+$",
        r"^修订时间\s*[：:]\s*\S+$",
        r"^\d{4}[-./]\s*\d{1,2}[-./]\s*\d{1,2}$",
        r"^版\s*次\s*[：:]\s*\S+$",
        r"^版本号\s*[：:]\s*\S+$",
        r"^(?:R\d{2,}|版本\s*\S+)\s*$",
        r"^修订次数\s*[：:]\s*\d+\s+\S+\s+页码",
        r"修订日期\s*[：:]\s*\d{4}-\d{1,2}-\d{1,2}\s*$",
        # 页码标记 -N-（公文底部页码）
        r"^-\d+-\s*$",
    ],
    # ========== 左 margin 编号列分离 ==========
    # 部分 PDF 排版将章节/条款编号放在页面左侧 margin 区域（如 x≈76），
    # 正文内容在右侧（如 x≈152）。pdfplumber 按 y 坐标聚行时会把两列混在一起，
    # 导致编号出现在行首或行尾不可预测，章节识别不稳定。
    # 开启后，x < margin_number_x 的 word 若匹配 margin_number_pattern，
    # 会被识别为编号并固定拼接到正文前面，保证 chapter_patterns 稳定匹配。
    # 设为 0 表示禁用此功能。
    "margin_number_x": 130,
    # 编号列中 word 必须匹配此正则才被视为"编号"（避免左侧其他内容被误识别）。
    # 默认匹配：多级数字编号（1.1, 1.1.2.1）和单字母编号（A, B, C）。
    "margin_number_pattern": r"^(?:\d+(?:\.\d+)*|[A-Z])$",
    # ========== 空白模板表格过滤 ==========
    # 空单元格率超过此阈值的表格视为"空白模板"（如签到表、申请表），
    # 将从解析结果中排除。这些表格无实际信息价值，会干扰检索和版本对比。
    # 设为 1.0 表示禁用此功能（不过滤任何表格）。
    "table_empty_cell_threshold": 0.6,
    # 这些字符（除 None 和纯空白外）也视为"空"单元格，
    # 与 table_empty_cell_threshold 配合使用。
    "table_empty_placeholders": ["□", "☐", "○", "——"],
    # 章节标题最大长度（超过此值的不视为标题）
    "max_chapter_title_length": 80,
    # 句末断句的最小段落长度（段落达到此长度且遇到句末终止符时断开）
    "sentence_break_min_length": 40,
}


def get_extract_config(config=None):
    """合并用户配置和默认配置"""
    result = copy.deepcopy(DEFAULT_CONFIG)
    if config and "extract" in config:
        result.update(config["extract"])
    return result


# ============================================================
# 智能后端选择
# ============================================================


def _quick_scan_pdf(filepath, sample_pages=5):
    """
    快速预扫描 PDF，收集决策所需的统计信息。
    只读取前 N 页（默认 5 页），耗时通常 < 0.5s。

    返回 dict:
      num_pages, sampled, avg_text_per_page, avg_tables_per_page,
      large_image_ratio, has_drawings_no_tables, text_samples
    """
    import pdfplumber

    result = {
        "num_pages": 0,
        "sampled": 0,
        "avg_text_per_page": 0,
        "avg_tables_per_page": 0,
        "large_image_ratio": 0.0,
        "has_drawings_no_tables": False,
        "text_samples": [],
    }

    try:
        with pdfplumber.open(filepath) as pdf:
            num_pages = len(pdf.pages)
            result["num_pages"] = num_pages
            if num_pages == 0:
                return result

            n = min(sample_pages, num_pages)
            total_text = 0
            total_tables = 0
            large_image_pages = 0
            drawings_no_tables_pages = 0

            for i in range(n):
                page = pdf.pages[i]
                text = page.extract_text() or ""
                tables = page.extract_tables()
                images = getattr(page, "images", [])
                # rects/lines/curves → 绘图对象
                rects = getattr(page, "rects", [])
                lines = getattr(page, "lines", [])
                curves = getattr(page, "curves", [])

                total_text += len(text)
                total_tables += len(tables)

                # 截取前 200 字符做样本
                if text:
                    result["text_samples"].append(text[:200])

                # 检测大图片（扫描件特征）
                page_area = page.width * page.height
                for img in images:
                    img_area = img.get("width", 0) * img.get("height", 0)
                    if page_area > 0 and img_area > page_area * 0.5:
                        large_image_pages += 1
                        break

                # 检测"有绘图线但无表格"（可能无框线表格被 pdfplumber 漏掉）
                has_drawings = len(rects) + len(lines) + len(curves) > 0
                if has_drawings and len(tables) == 0:
                    drawings_no_tables_pages += 1

            result["sampled"] = n
            result["avg_text_per_page"] = total_text / n
            result["avg_tables_per_page"] = total_tables / n
            result["large_image_ratio"] = large_image_pages / n
            result["has_drawings_no_tables"] = drawings_no_tables_pages > n * 0.4

    except Exception as e:
        logger.debug(f"预扫描失败: {e}")

    return result


def _detect_borderless_table_hint(scan):
    """
    基于文本样本检测无框线表格的线索。

    判断依据：
    - 文本中出现表头关键词（序号、名称、描述/风险/措施…）
    - 短行密集且包含连续空格或制表符（列对齐特征）
    """
    table_keywords = [
        "序号",
        "名称",
        "描述",
        "风险",
        "措施",
        "类别",
        "编号",
        "责任人",
        "频率",
        "要求",
        "备注",
        "检查项",
        "标准",
    ]
    keyword_hits = 0
    aligned_line_hits = 0

    for sample in scan.get("text_samples", []):
        # 关键词命中
        for kw in table_keywords:
            if kw in sample:
                keyword_hits += 1
                break
        # 连续空格/制表符 → 列对齐
        if re.search(r"\S+\s{3,}\S+\s{3,}\S+", sample):
            aligned_line_hits += 1

    total = max(1, len(scan.get("text_samples", [])))
    return (keyword_hits / total >= 0.5) and (aligned_line_hits / total >= 0.3)


def _select_backend(filepath, cfg):
    """
    智能选择解析后端。

    决策流程：
    1. 扫描件检测  → 文本量极低 或 大图片覆盖页面 → MinerU
    2. 无框线表格  → 有绘图对象但 pdfplumber 提取不到表格 → MinerU
    3. 文本表格线索 → 文本中出现表头关键词且列对齐 → MinerU
    4. 正常文档    → pdfplumber

    返回 (backend_name, reason)
    """
    scan = _quick_scan_pdf(filepath)

    if scan["num_pages"] == 0:
        return "pdfplumber", "空 PDF"

    # 1. 扫描件：平均每页文字 < 50 字符
    if scan["avg_text_per_page"] < 50:
        return "mineru", f"疑似扫描件（平均 {scan['avg_text_per_page']:.0f} 字/页）"

    # 2. 大图片覆盖 > 50% 采样页
    if scan["large_image_ratio"] > 0.5:
        return "mineru", f"疑似扫描件（大图片覆盖率 {scan['large_image_ratio']:.0%}）"

    # 3. 有绘图线但 pdfplumber 提取不到表格 → 可能无框线表格
    if scan["has_drawings_no_tables"] and scan["avg_tables_per_page"] < 0.3:
        return "mineru", "检测到绘图对象但 pdfplumber 未提取到表格（疑似无框线表格）"

    # 4. 文本中出现表头关键词 + 列对齐特征
    if _detect_borderless_table_hint(scan) and scan["avg_tables_per_page"] < 0.3:
        return "mineru", "文本中出现表格关键词及列对齐特征（疑似无框线表格）"

    # 5. pdfplumber 足以应对
    if scan["avg_tables_per_page"] >= 1:
        return "pdfplumber", f"表格提取正常（平均 {scan['avg_tables_per_page']:.1f} 表/页）"

    return "pdfplumber", "文档特征正常"


# ============================================================
# PDF 解析（通用版）
# ============================================================


def extract_pdf(filepath, config=None):
    """
    通用 PDF 解析

    策略（单次打开 PDF，两遍遍历）：
    第一遍：收集每页文字行 + 提取表格 + 获取表格区域坐标
    → 统计高频重复行
    第二遍：排除表格区域和重复行，拼接正文流
    → 在正文流上按语义分段 + 反查页码

    后端选择：
    config["extract"]["backend"] = "mineru" → 使用 MinerU VLM/OCR 引擎
    config["extract"]["backend"] = "auto"   → 智能选择（预扫描后决定）
    默认使用 pdfplumber
    """
    cfg = get_extract_config(config)

    # 后端选择
    backend = cfg.get("backend", "pdfplumber")
    if backend == "mineru":
        from doc_parser.mineru_backend import extract_pdf_with_mineru

        return extract_pdf_with_mineru(filepath, config)
    elif backend == "auto":
        chosen, reason = _select_backend(filepath, cfg)
        if chosen == "mineru":
            try:
                from doc_parser.mineru_backend import extract_pdf_with_mineru

                logger.info(f"自动选择 MinerU 后端：{reason}")
                return extract_pdf_with_mineru(filepath, config)
            except RuntimeError:
                # MinerU 未安装，降级到 pdfplumber
                logger.warning(f"{reason}，但 MinerU 未安装，降级到 pdfplumber")
        else:
            if reason:
                logger.info(f"自动选择 pdfplumber 后端：{reason}")

    with pdfplumber.open(filepath) as pdf:
        num_pages = len(pdf.pages)
        if num_pages == 0:
            return Document(filename=os.path.basename(filepath), paragraphs=[], tables=[])

        # 编译 margin 编号正则（一次编译，全页复用）
        margin_number_x = cfg.get("margin_number_x", 0)
        margin_number_re = None
        if margin_number_x > 0:
            margin_number_re = re.compile(cfg.get("margin_number_pattern", r"^(?:\d+(?:\.\d+)*|[A-Z])$"))

        # ========== 第一遍：收集行 + 提取表格 + 获取表格区域 ==========
        all_lines_flat = []  # 所有行的平面列表（用于统计重复）
        page_data = []  # [(page_num, page_lines, table_regions)]
        tables = []
        current_chapter = ""
        current_chapter_title = ""

        for page_num, page in enumerate(pdf.pages, 1):
            page_height = page.height

            # 计算页眉/页脚的 Y 坐标边界
            header_y = page_height * (cfg["header_margin_pct"] / 100)
            footer_y = page_height * (1 - cfg["footer_margin_pct"] / 100)

            # 获取文字对象（带坐标）
            words = page.extract_words(keep_blank_chars=False, use_text_flow=True)

            # 按行聚合（按 top 坐标分组，容差 3pt）
            page_lines = _words_to_lines(
                words, y_tolerance=3, margin_number_x=margin_number_x, margin_number_re=margin_number_re
            )

            # 收集有效行用于重复统计
            for line_top, line_text in page_lines:
                if line_top < header_y or line_top > footer_y:
                    continue
                all_lines_flat.append(line_text.strip())

            # 提取表格 + 获取表格区域坐标
            page_tables = page.find_tables()
            table_regions = [(t.bbox[1], t.bbox[3]) for t in page_tables]  # (top, bottom)

            # 从页面文字中更新章节追踪
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    ch = _detect_chapter(line.strip(), cfg)
                    if ch:
                        current_chapter, current_chapter_title = ch

            # 提取表格数据
            for pt in page_tables:
                table_data = pt.extract()
                if table_data and len(table_data) >= 2:
                    cleaned = _clean_table(table_data)
                    if cleaned and len(cleaned) >= 2:
                        tables.append(
                            Table(
                                rows=cleaned,
                                page=page_num,
                                chapter=current_chapter,
                                chapter_title=current_chapter_title,
                                source_file=os.path.basename(filepath),
                                index=len(tables) + 1,
                            )
                        )

            page_data.append((page_num, page_lines, table_regions, header_y, footer_y))

        # ========== 统计高频重复行（自动识别页眉页脚残留）==========
        repeat_threshold = max(3, int(num_pages * cfg["repeat_line_threshold_pct"] / 100))
        line_counts = Counter(all_lines_flat)
        repeated_lines = {line for line, count in line_counts.items() if count >= repeat_threshold and len(line) < 100}

        # ========== 第二遍：构建正文流（排除表格区域 + 重复行）==========
        full_text = ""
        page_boundaries = []  # [(start_offset, end_offset, page_num)]

        for page_num, page_lines, table_regions, header_y, footer_y in page_data:
            start_offset = len(full_text)
            for line_top, line_text in page_lines:
                # 跳过页眉/页脚
                if line_top < header_y or line_top > footer_y:
                    continue
                # 跳过重复行
                if line_text.strip() in repeated_lines:
                    continue
                # 跳过表格区域内的行
                if _is_in_table_region(line_top, table_regions):
                    continue
                full_text += line_text + "\n"

            end_offset = len(full_text)
            if end_offset > start_offset:
                page_boundaries.append((start_offset, end_offset, page_num))

    # ========== 跨页表格合并 (启发式) ==========
    tables = _merge_cross_page_tables(tables)

    # ========== 空白模板表格过滤 ==========
    tables = _filter_template_tables(tables, cfg)

    # ========== 流式分段 + 章节标记 + 反查页码 ==========
    paragraphs = _segment_and_locate(full_text, page_boundaries, filepath, cfg)

    # 为表格分配章节信息
    _assign_table_chapters(tables, paragraphs)

    return Document(filename=os.path.basename(filepath), paragraphs=paragraphs, tables=tables)


def _words_to_lines(words, y_tolerance=3, margin_number_x=0, margin_number_re=None):
    """
    将 pdfplumber 的 word 对象按 Y 坐标聚合为行
    返回: [(line_top_y, line_text), ...]

    当 margin_number_x > 0 时，启用编号列分离：
    - x0 < margin_number_x 且文本匹配 margin_number_re 的 word → 视为编号
    - 同一行内，输出格式固定为 "编号 正文"（编号在前）
    - 保证 chapter_patterns（^\\d+\\.\\d+ xxx）能稳定匹配
    """
    if not words:
        return []

    # 按 top 排序
    sorted_words = sorted(words, key=lambda w: (w["top"], w["x0"]))

    lines = []
    current_line_words = [sorted_words[0]]
    current_top = sorted_words[0]["top"]

    for word in sorted_words[1:]:
        if abs(word["top"] - current_top) <= y_tolerance:
            current_line_words.append(word)
        else:
            # 新行
            line_text = _assemble_line(current_line_words, margin_number_x, margin_number_re)
            lines.append((current_top, line_text))
            current_line_words = [word]
            current_top = word["top"]

    # 最后一行
    if current_line_words:
        line_text = _assemble_line(current_line_words, margin_number_x, margin_number_re)
        lines.append((current_top, line_text))

    return lines


def _assemble_line(line_words, margin_number_x, margin_number_re):
    """
    将同一行的 word 列表拼接为文本。

    如果启用了 margin 编号列（margin_number_x > 0），则分离编号和正文，
    统一输出为 "编号 正文"，消除 PDF 文字流顺序的不确定性。
    """
    sorted_by_x = sorted(line_words, key=lambda w: w["x0"])

    if margin_number_x <= 0 or margin_number_re is None:
        # 未启用，保持原行为：按 x 坐标顺序拼接
        return " ".join(w["text"] for w in sorted_by_x)

    # 分离编号列和正文列
    number_parts = []
    body_parts = []
    for w in sorted_by_x:
        if w["x0"] < margin_number_x and margin_number_re.match(w["text"].strip()):
            number_parts.append(w["text"].strip())
        else:
            body_parts.append(w["text"])

    body_text = " ".join(body_parts)

    if number_parts:
        # 多段编号片段（极少见）用 . 连接；通常只有一个
        number_text = ".".join(number_parts) if len(number_parts) > 1 else number_parts[0]
        return f"{number_text} {body_text}" if body_text else number_text
    else:
        return body_text


def _is_in_table_region(line_top, table_regions):
    """判断某行是否在表格区域内"""
    return any(table_top <= line_top <= table_bottom for table_top, table_bottom in table_regions)


def _segment_and_locate(full_text, page_boundaries, filepath, cfg):
    """在正文流上按语义分段，标注页码和章节"""
    if not full_text.strip():
        return []

    min_len = cfg["min_paragraph_length"]
    noise_patterns = [re.compile(p) for p in cfg.get("noise_patterns", [])]

    # 分段
    raw_paras = _split_stream(full_text, cfg)

    # 组装 Paragraph 对象
    paragraphs = []
    current_chapter = ""
    current_chapter_title = ""

    for para_text, char_start, char_end in raw_paras:
        para_text = para_text.strip()

        # 章节检测（先检测，标题段免长度过滤）
        ch = _detect_chapter(para_text, cfg)

        # 长度过滤（章节标题段免过滤，因为 "2 日常管理" 等标题较短）
        if not ch and len(para_text) < min_len:
            continue

        # 噪声正则过滤
        if any(p.match(para_text) for p in noise_patterns):
            continue

        if ch:
            current_chapter, current_chapter_title = ch

        # 反查页码
        page_start = _find_page(char_start, page_boundaries)
        page_end = _find_page(max(char_end - 1, char_start), page_boundaries)

        paragraphs.append(
            Paragraph(
                text=para_text,
                page=page_start,
                page_end=page_end if page_end != page_start else 0,
                chapter=current_chapter,
                chapter_title=current_chapter_title,
                source_file=os.path.basename(filepath),
                index=len(paragraphs) + 1,
            )
        )

    return paragraphs


def _split_stream(full_text, cfg):
    """
    在流式全文上按语义分段

    分割信号（优先级从高到低）：
    1. 章节标题行前断开
    2. 连续空行
    3. 句末终止符 + 换行（段落已达一定长度）
    4. 长度上限强制断开
    """
    max_len = cfg["max_paragraph_length"]
    noise_line_re = [re.compile(p) for p in cfg.get("noise_line_patterns", [])]
    paragraphs = []
    lines = full_text.split("\n")
    current_lines = []
    current_start = 0
    char_pos = 0

    for line in lines:
        line_start = char_pos
        char_pos += len(line) + 1  # +1 for \n

        stripped = line.strip()

        # 空行 → 断开
        if not stripped:
            if current_lines:
                para_text = " ".join(current_lines)
                if len(para_text.strip()) > 0:
                    paragraphs.append((para_text, current_start, line_start))
                current_lines = []
            current_start = char_pos
            continue

        # 元数据行剥离：整行匹配配置的版本管理元数据模式 → 跳过（不进正文段落）
        # 避免"修订日期：2026-05-08"这类独立行被并入相邻正文造成句子拼接。
        if noise_line_re and any(p.match(stripped) for p in noise_line_re):
            if current_lines:
                para_text = " ".join(current_lines)
                if len(para_text.strip()) > 0:
                    paragraphs.append((para_text, current_start, line_start))
                current_lines = []
            current_start = char_pos
            continue

        # 章节标题 → 独占一段（先断开前面的，标题自身也单独成段）
        if _detect_chapter(stripped, cfg):
            if current_lines:
                para_text = " ".join(current_lines)
                if len(para_text.strip()) > 0:
                    paragraphs.append((para_text, current_start, line_start))
                current_lines = []
            # 标题独占一段
            paragraphs.append((stripped, line_start, char_pos))
            current_start = char_pos
            continue

        # 正文行 → 先累积到 current_lines
        current_lines.append(stripped)
        current_text = " ".join(current_lines)

        # 句末终止符 + 段落已有一定长度 → 断开
        if (
            stripped
            and stripped[-1] in "。；！？.;!?"
            and len(current_text) >= cfg.get("sentence_break_min_length", 40)
        ) or (len(current_text) > max_len and stripped and stripped[-1] in "。.;；，,"):
            paragraphs.append((current_text, current_start, char_pos))
            current_lines = []
            current_start = char_pos

    # 尾部
    if current_lines:
        para_text = " ".join(current_lines)
        if len(para_text.strip()) > 0:
            paragraphs.append((para_text, current_start, char_pos))

    return paragraphs


# 列表项特征：以动作动词开头（后面通常跟长描述）
_LIST_VERB_PREFIXES = (
    "负责",
    "建立",
    "整合",
    "加强",
    "完成",
    "管理与",
    "参与",
    "组织",
    "规划",
    "做好",
    "开展",
    "制定",
    "依据",
    "按照",
    "定期",
    "管理维护",
    "采集",
    "编制",
    "审批",
    "评估",
)

# 列表项结尾标点（真实章节标题不会以这些结尾）
_LIST_END_PUNCT = "；，。、；,.；"


# 单纯数字编号的正则（用于在 _detect_chapter 中识别并施加更严格的过滤）
_SINGLE_NUMBER_RE = re.compile(r"^\d+\s+(.+)")

# 中文编号正则（用于判断是否为中文数字编号模式，施加标题终止符截断）
_CHINESE_NUM_RE = re.compile(r"^（?[一二三四五六七八九十]+[、）)]")

# 标题终止符：标题后紧跟正文时，标题通常以这些符号结尾
# 遇到这些符号且后面还有内容时，只取终止符前的部分作为标题
_HEADING_TERMINATORS = ("．", "。", "：", ":")


def _detect_chapter(text, cfg):
    """识别章节标题

    过滤规则（避免表格行/目录条目/正文碎片/列表项被误识别为标题）：
    - 标题长度不超过 max_chapter_title_length
    - 标题部分长度 >= 2（排除 "6 R" 之类的单字符）
    - 标题不含日期模式（排除 "3 2025-12-03 ..." 之类的表格行）
    - 标题不含过多列表分隔符（排除 "0.2-1、0.4-1、0.4-6、 ..." 之类的内容）
    - 标题不以列表项标点结尾（排除 "...负责...工作；" 之类的列表项）
    - 标题不以动作动词开头且过长（排除 "负责信息系统建设..." 之类的列表项）
    - 单纯数字编号 ^\\d+ 有额外限制：标题长度和编号大小（见配置项）
    - 中文编号标题含终止符时截断标题（处理"标题+正文同行"的情况）
    """
    text = text.strip()
    if not text or len(text) > cfg.get("max_chapter_title_length", 80):
        return None
    for pattern in cfg["chapter_patterns"]:
        m = re.match(pattern, text)
        if m:
            chapter = m.group(1)
            title = m.group(2).strip()
            # 中文编号：标题含终止符时截断（处理 "（一）恶劣天气运行风险． 7-8 月..." 的情况）
            if _CHINESE_NUM_RE.match(text):
                for term in _HEADING_TERMINATORS:
                    idx = title.find(term)
                    if 0 < idx < len(title) - 1:
                        title = title[:idx].strip()
                        break
            # 标题太短 → 可能是表格数字
            if len(title) < 2:
                return None
            # 标题含日期 → 可能是表格行
            if re.search(r"\d{4}[-./]\d{1,2}[-./]\d{1,2}", title):
                return None
            # 标题含过多列表分隔符 → 可能是正文碎片
            if title.count("、") + title.count("，") > 2:
                return None
            # 标题以列表项标点结尾 → 列表项，非章节标题
            if title[-1] in _LIST_END_PUNCT:
                return None
            # 标题以动作动词开头且较长 → 列表项描述，非章节标题
            if len(title) > 12 and title.startswith(_LIST_VERB_PREFIXES):
                return None
            # 单纯数字编号 ^\d+ 的额外过滤：
            # - 编号值过大（如 2026）→ 可能是年份，非章节号
            # - 标题过长 → 可能是正文行误匹配（如 "6 月份，管理局通过..."）
            if _SINGLE_NUMBER_RE.match(text) and not re.match(r"^\d+\.\d+", text):
                try:
                    num_val = int(chapter)
                except ValueError:
                    num_val = 0
                if num_val > cfg.get("single_number_max_value", 999):
                    return None
                if len(title) > cfg.get("single_number_title_max_length", 20):
                    return None
            return (chapter, title)
    return None


def _find_page(offset, page_boundaries):
    """根据字符偏移反查页码"""
    for start, end, page_num in page_boundaries:
        if start <= offset < end:
            return page_num
    return page_boundaries[-1][2] if page_boundaries else 0


def _clean_table(table_data):
    """清理表格数据"""
    cleaned = [row for row in table_data if any(cell and cell.strip() for cell in row)]
    if not cleaned:
        return []
    cleaned = [[cell.strip() if cell else "" for cell in row] for row in cleaned]
    num_cols = len(cleaned[0])
    non_empty_cols = [c for c in range(num_cols) if any(row[c] if c < len(row) else "" for row in cleaned)]
    if non_empty_cols:
        cleaned = [[row[c] if c < len(row) else "" for c in non_empty_cols] for row in cleaned]
    return cleaned


def _filter_template_tables(tables, cfg):
    """
    过滤空白模板表格（签到表/申请表等无信息价值的空表格）。

    判定条件：空单元格率超过 table_empty_cell_threshold。
    空单元格包括：None、纯空白、以及 table_empty_placeholders 中定义的占位符。
    """
    threshold = cfg.get("table_empty_cell_threshold", 1.0)
    if threshold >= 1.0:
        return tables  # 功能禁用

    placeholders = set(cfg.get("table_empty_placeholders", []))

    result = []
    for table in tables:
        total_cells = sum(len(row) for row in table.rows)
        if total_cells == 0:
            continue
        empty_count = 0
        for row in table.rows:
            for cell in row:
                if not cell or not cell.strip() or cell.strip() in placeholders:
                    empty_count += 1
        if empty_count / total_cells < threshold:
            result.append(table)
    # 重新标号
    for idx, t in enumerate(result, 1):
        t.index = idx
    return result


def _assign_table_chapters(tables, paragraphs):
    """为表格分配最近的章节信息"""
    if not paragraphs:
        return
    for table in tables:
        if table.chapter:
            continue
        best = None
        for p in paragraphs:
            if p.page <= table.page and p.chapter:
                best = p
        if best:
            table.chapter = best.chapter
            table.chapter_title = best.chapter_title


def _merge_cross_page_tables(tables: list) -> list:
    """
    启发式合并跨页连续表格。

    合并条件（同时满足）:
    1. 同一 source_file
    2. 页码连续 (后续表格的 page == 前一表格的 page/page_end + 1)
    3. 列数接近 (差异 ≤ 1)
    续页的表头相似度仅用于 _append 阶段决定是否跳过重复行，不影响是否合并。

    合并后:
    - 保留首张表格的 headers / page / chapter
    - 续页的数据行追加到首张的 rows 末尾（跳过重复表头行）
    - 更新 page_end = 末页页码
    """
    if len(tables) <= 1:
        return tables

    merged: list = []
    i = 0
    while i < len(tables):
        cur = tables[i]
        j = i + 1
        while j < len(tables):
            nxt = tables[j]
            if not _should_merge_tables(cur, nxt):
                break
            # 执行合并: cur.rows += nxt 的非表头行
            new_rows = _append_rows_skip_dup_header(cur, nxt)
            cur.rows = new_rows
            # 更新页码范围
            cur.page_end = getattr(nxt, "page_end", None) or nxt.page
            # 继承章节（取第一个有的）
            if not cur.chapter and getattr(nxt, "chapter", None):
                cur.chapter = nxt.chapter
                cur.chapter_title = nxt.chapter_title
            j += 1
        merged.append(cur)
        i = max(i + 1, j)

    # 重新标号
    for idx, t in enumerate(merged, 1):
        t.index = idx
    return merged


def _should_merge_tables(a, b, header_threshold: float = 0.85) -> bool:
    """判断两张表格是否应为同一逻辑表格的跨页片段。

    合并判定: 同文件 + 页码连续 + 列数兼容 (差异 ≤ 1) + 重复表头。

    重复表头是识别“续表”的关键证据。仅按相邻页和列数合并会把两张
    恰好列数相同的独立表格拼在一起；没有可靠续表证据时，宁可保留两表。
    """
    # 1. 同文件
    if a.source_file != b.source_file:
        return False

    # 2. 页码连续
    a_end = a.page_end or a.page
    if b.page != a_end + 1:
        return False

    # 3. 列数接近
    a_rows = a.rows or []
    b_rows = b.rows or []
    if not a_rows or not b_rows:
        return False
    a_cols = len(a_rows[0]) if a_rows[0] else 0
    b_cols = len(b_rows[0]) if b_rows[0] else 0
    if abs(a_cols - b_cols) > 1:
        return False
    return _row_token_jaccard(a_rows[0], b_rows[0]) >= header_threshold


def _row_token_jaccard(row_a, row_b) -> float:
    """两行之间的 Token Jaccard 相似度。"""
    tokens_a = {str(c).strip() for c in row_a if c and str(c).strip()}
    tokens_b = {str(c).strip() for c in row_b if c and str(c).strip()}
    if not tokens_a or not tokens_b:
        return 0.0
    return len(tokens_a & tokens_b) / len(tokens_a | tokens_b)


def _append_rows_skip_dup_header(target, source, header_threshold: float = 0.85) -> list:
    """追加 source 的行到 target 末尾，跳过重复的表头行（如果有）。"""
    result = list(target.rows or [])
    source_rows = list(source.rows or [])
    if not source_rows:
        return result

    target_header = result[0] if result else None
    target_cols = len(target_header) if target_header else len(source_rows[0])

    for idx, row in enumerate(source_rows):
        # 续页第 0 行若与表头相似 → 跳过（重复表头）
        if idx == 0 and target_header and _row_token_jaccard(target_header, row) >= header_threshold:
            continue
        # 列数对齐: 短的补空, 长的截断
        if len(row) < target_cols:
            row = list(row) + [""] * (target_cols - len(row))
        elif len(row) > target_cols:
            row = row[:target_cols]
        result.append(row)
    return result


# ============================================================
# Word 解析
# ============================================================


def _iter_docx_blocks(doc):
    """按 OOXML body 顺序产出顶层段落和表格。"""
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", DocxParagraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield "table", DocxTable(child, doc)


def extract_docx(filepath, config=None):
    """从 Word 文档提取段落和表格"""
    cfg = get_extract_config(config)
    doc = DocxDocument(filepath)
    paragraphs = []
    tables = []
    current_chapter = ""
    current_chapter_title = ""

    # 必须按 body 块顺序读取：doc.paragraphs 和 doc.tables 是两份独立列表，
    # 分别遍历会让所有表格错误地继承最后一个章节。
    for block_order, (block_type, block) in enumerate(_iter_docx_blocks(doc), 1):
        if block_type == "paragraph":
            text = block.text.strip()
            if not text:
                continue

            # 与 PDF 路径一致：先识别标题，短标题不受最小段落长度过滤。
            ch = _detect_chapter(text, cfg)
            if not ch and len(text) < cfg["min_paragraph_length"]:
                continue
            if ch:
                current_chapter, current_chapter_title = ch

            paragraphs.append(
                Paragraph(
                    text=text,
                    page=0,
                    chapter=current_chapter,
                    chapter_title=current_chapter_title,
                    source_file=os.path.basename(filepath),
                    index=len(paragraphs) + 1,
                    order=block_order,
                )
            )
            continue

        rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
        cleaned = _clean_table(rows)
        if cleaned and len(cleaned) >= 2:
            tables.append(
                Table(
                    rows=cleaned,
                    page=0,
                    chapter=current_chapter,
                    chapter_title=current_chapter_title,
                    context_before="",
                    source_file=os.path.basename(filepath),
                    index=len(tables) + 1,
                    order=block_order,
                )
            )

    _assign_table_chapters(tables, paragraphs)

    # 过滤空白模板表格（与 PDF 路径保持一致）
    tables = _filter_template_tables(tables, cfg)

    return Document(filename=os.path.basename(filepath), paragraphs=paragraphs, tables=tables)


# ============================================================
# 统一入口
# ============================================================


def extract_document(filepath, config=None):
    """统一文档解析入口"""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(filepath, config)
    elif ext == ".docx":
        return extract_docx(filepath, config)
    elif ext == ".doc":
        raise ValueError("不支持的格式: .doc（旧版 Word 二进制格式）。请先转换为 .docx 格式。")
    else:
        raise ValueError(f"不支持的格式: {ext}")


# ============================================================
# 公共 API
# ============================================================


def parse(filepath: str, config: dict | None = None) -> Document:
    """
    解析文档为结构化段落+表格（公共 API）

    Args:
        filepath: 文件路径 (支持 .pdf, .docx)
        config: 解析配置字典，可选。支持的 key:
            - header_margin_pct: 页眉区域百分比 (默认 8)
            - footer_margin_pct: 页脚区域百分比 (默认 8)
            - repeat_line_threshold_pct: 重复行检测阈值 (默认 30)
            - min_paragraph_length: 最小段落长度 (默认 10)
            - max_paragraph_length: 最大段落长度 (默认 600)

    Returns:
        Document 对象，包含 .paragraphs 和 .tables

    Example:
        from doc_parser import parse
        doc = parse("manual.pdf", config={"min_paragraph_length": 20})
        for para in doc.paragraphs:
            print(para.text, para.location)
    """
    return extract_document(filepath, config)


def parse_to_markdown(filepath: str, config: dict | None = None) -> str:
    """
    解析文档并直接转为 Markdown（便捷 API）

    等价于 parse(filepath, config).to_markdown()

    Args:
        filepath: 文件路径 (支持 .pdf, .docx)
        config: 解析配置字典，可选

    Returns:
        Markdown 格式的字符串

    Example:
        from doc_parser import parse_to_markdown
        md = parse_to_markdown("manual.pdf")
        with open("manual.md", "w", encoding="utf-8") as f:
            f.write(md)
    """
    return parse(filepath, config).to_markdown()
