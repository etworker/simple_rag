"""Word (.docx) 解析实现"""

import os

from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from doc_parser._tables import (
    assign_table_chapters,
    clean_table,
    filter_template_tables,
)
from doc_parser._text import detect_chapter
from doc_parser.models import Document, Paragraph, Table


def _iter_docx_blocks(doc):
    """按 OOXML body 顺序产出顶层段落和表格。"""
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", DocxParagraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield "table", DocxTable(child, doc)


def extract_docx(filepath, config=None, get_config=None):
    """从 Word 文档提取段落和表格"""
    if get_config is None:
        from doc_parser.parser import get_extract_config

        get_config = get_extract_config
    cfg = get_config(config)
    doc = DocxDocument(filepath)
    paragraphs = []
    tables = []
    current_chapter = ""
    current_chapter_title = ""

    # 必须按 body 块顺序读取：doc.paragraphs 和 doc.tables 是两份独立列表，
    # 分别遍历会让所有表格错误地继承最后一个章节。
    for block_order, (block_type, block) in enumerate(_iter_docx_blocks(doc), 1):
        if block_type == "paragraph":
            text = block.text.strip()
            if not text:
                continue

            # 与 PDF 路径一致：先识别标题，短标题不受最小段落长度过滤。
            ch = detect_chapter(text, cfg)
            if not ch and len(text) < cfg["min_paragraph_length"]:
                continue
            if ch:
                current_chapter, current_chapter_title = ch

            paragraphs.append(
                Paragraph(
                    text=text,
                    page=0,
                    chapter=current_chapter,
                    chapter_title=current_chapter_title,
                    source_file=os.path.basename(filepath),
                    index=len(paragraphs) + 1,
                    order=block_order,
                )
            )
            continue

        rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
        cleaned = clean_table(rows)
        if cleaned and len(cleaned) >= 2:
            tables.append(
                Table(
                    rows=cleaned,
                    page=0,
                    chapter=current_chapter,
                    chapter_title=current_chapter_title,
                    context_before="",
                    source_file=os.path.basename(filepath),
                    index=len(tables) + 1,
                    order=block_order,
                )
            )

    assign_table_chapters(tables, paragraphs)

    # 过滤空白模板表格（与 PDF 路径保持一致）
    tables = filter_template_tables(tables, cfg)

    return Document(filename=os.path.basename(filepath), paragraphs=paragraphs, tables=tables)
